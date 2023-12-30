#!/usr/bin/env python3
import json
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

file_name = "" #path to file

severity_to_num = {
        "Critical": 5,
        "High": 4,
        "Moderate": 3,
        "Low": 2,
        "Infomational": 1
    }

num_to_severity = {
        5: "Critical",
        4: "High",
        3: "Medium",
        2: "Low",
        1: "Infomational"
    }

severity_to_color = {
    5: "ec421f",
    4: "f06d53",
    3: "f6b26b",
    2: "ffe599",
    1: "a4c2f4"
}

def extract_values(json_report):
    alerts = []

    for alert in json_report:
        description = alert["rule_data"]["description"]
        severity = severity_to_num[alert["rule_data"]["severity"]]
        risk_categories = alert["rule_data"]["risk_categories"] if (alert["rule_data"]["risk_categories"]) else ""
        rule_category = alert["rule_data"]["rule_category"]
        resources = []
        
        for alert_instance in alert["alert_instances"]:
            if(alert_instance["status"] != "passed"):
                resources.append(alert_instance["resource_id"])
            else:
                continue

        if(resources):
            alerts.append(
                {
                    'Description': description,
                    'Severity': severity,
                    'Category': rule_category,
                    'Risks': ", ".join(risk_categories) if(len(risk_categories)>0) else risk_categories,
                    'resources': "- " + '\n- '.join(resources)
                    }
                ) 
        else:
            continue
    
    return alerts

def create_custom_table(doc:docx.Document, data:list):
    """
    Populate document with the table created from the dictory data 
    Inputs:
        doc = docx.Document
        data = dictionary
        return = table created
    
    """
    table = doc.add_table(rows=len(data)+1, cols=2)

    # table.rows[0].height = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Merge the cells in the first row
    first_row = table.rows[0]
    first_row.cells[0].merge(first_row.cells[1])

    table.allow_autofit = True

    table.columns[0].width = Inches(1.2)
    table.columns[1].width = doc.sections[0].page_width - Inches(3.2)

    color_cell_by_severity(table.cell(0,0), data["Severity"])
    
    # set the size and color of the border
    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            set_cell_border(table.cell(i, j))

    data["Severity"] = num_to_severity[data['Severity']]

    table.cell(0,0).text = data['Severity'] + ": " + data["Description"]
    for i, key in enumerate(data.keys(), 1):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = data[key]

    return table


def bold_column(table:docx.table.Table, column:int = 0):
    """Function to make specified column bold."""
    for i in range(len(table.rows)):
        cell = table.cell(i, column)
        for p in cell.paragraphs:
            p.runs[0].bold = True
    return

def color_cell_by_severity(cell, severity):
    color = severity_to_color[severity]

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    
    cell._tc.get_or_add_tcPr().append(shading_elm)
    return


def set_cell_border(cell, size=4, color="000000"):
    """
    Set border with size and color specified
    
    """
    cell_tcPr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')

    for border_type in ['top', 'left', 'bottom', 'right']:
        border_elm = OxmlElement(f'w:{border_type}')
        border_elm.set(qn('w:val'), 'single')
        border_elm.set(qn('w:sz'), str(size))
        border_elm.set(qn('w:space'), '0')
        border_elm.set(qn('w:color'), color)
        tc_borders.append(border_elm)

    cell_tcPr.append(tc_borders)
    return

def setup_page(doc):
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.left_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    format = style.paragraph_format
    format.space_before = Pt(4)
    format.space_after = Pt(4)

    return

def main():
    try:
        with open(file_name) as f:
            report = f.read()
    except FileNotFoundError:
        print("File cannot found. Please make sure the file path and name are correct.")
        exit(1)

    doc = docx.Document()

    setup_page(doc)

    alerts = extract_values(json.loads(report))
    
    sorted_alerts = sorted(alerts, key=lambda x:x["Severity"], reverse=True)

    for alert in sorted_alerts:
        table = create_custom_table(doc, alert)
        bold_column(table)
        doc.add_paragraph()

    doc.save("alerts.docx")

if __name__ == "__main__":
    main()
    